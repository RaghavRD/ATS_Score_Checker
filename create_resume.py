from docx import Document

doc = Document()
doc.add_heading('John Doe', 0)
doc.add_paragraph('Python Developer with 5 years of experience in Django and AI.')
doc.add_heading('Experience', level=1)
doc.add_paragraph('Senior Developer at Tech Corp. Built ATS systems.')
doc.add_heading('Education', level=1)
doc.add_paragraph('B.Sc. in Computer Science.')
doc.save('test_resume.docx')
print("Created test_resume.docx")
